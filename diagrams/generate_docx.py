import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from PIL import Image

doc = docx.Document()

figures = [
    ("fig1_system_architecture.png", "Fig. 1"),
    ("fig2_address_resolution.png", "Fig. 2"),
    ("fig3_escrow_payment.png", "Fig. 3"),
    ("fig4_courier_matching.png", "Fig. 4"),
    ("fig5_offline_architecture.png", "Fig. 5"),
    ("fig6_agent_network.png", "Fig. 6")
]

total = len(figures)

for i, (filename, title) in enumerate(figures):
    if i == 0:
        section = doc.sections[0]
    else:
        # Create a new page with a new section block to allow for unique footers
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        
    # Apply precise margins
    section.top_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.bottom_margin = Cm(1.0)
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)

    # Detach headers and footers from previous sections to keep them unique
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    # Write Sheet Number natively into the Document's Header
    header = section.header
    p_head = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p_head.text = "" # Clear default
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_head = p_head.add_run(f"{i+1}/{total}")
    run_head.font.name = 'Arial'
    run_head.font.size = Pt(12)

    # Write Figure Number natively into the Document's Footer
    footer = section.footer
    p_foot = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_foot.text = "" # Clear default
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run(title)
    run_foot.font.name = 'Arial'
    run_foot.font.size = Pt(12)

    # Add the Image to the page body
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    
    try:
        img = Image.open(filename)
        w, h = img.size
        ratio = w / h
        
        # Give ample room so the image vertically sits between the header and footer comfortably
        max_w_cm = 16.5
        max_h_cm = 21.0
        
        new_w_cm = max_w_cm
        new_h_cm = new_w_cm / ratio
        
        if new_h_cm > max_h_cm:
            new_h_cm = max_h_cm
            new_w_cm = new_h_cm * ratio
            
        run_img.add_picture(filename, width=Cm(new_w_cm), height=Cm(new_h_cm))
    except Exception as e:
        print(f"Error loading {filename}: {e}")

doc.save("KIPI_Drawings.docx")
print("KIPI_Drawings.docx rewritten using Sections for pinned Headers and Footers.")
